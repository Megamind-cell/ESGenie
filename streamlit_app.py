import streamlit as st
import numpy as np
import faiss
import json
import openai
import tiktoken
import pandas as pd
from PIL import Image
from collections import defaultdict
import os
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import re
from io import BytesIO
from docx import Document
from datetime import datetime

# ====== Streamlit Config ======
st.set_page_config(page_title=" RFP Assistant ", layout="wide")

# ====== Sidebar Instructions ======
st.sidebar.title("📘 How to Use")

st.sidebar.markdown("""
### Ask a Question  
Paste an RFP prompt or type a plain-language query.

### Review the Draft  
Each answer comes with inline citations, verbatim excerpts and live document links—click to open the source.

### Refine & Export  
Copy/export the response into your RFP template, then tailor with client-specific nuance and tone adjustments.
""")

st.sidebar.markdown("---")

st.sidebar.markdown("""
### 💡 Example Questions  
- “Outline Bain’s net-zero emissions commitment and progress to date.”  
- “Which ESG policies apply to Bain’s suppliers?”  
- “Provide diversity metrics for Bain’s global workforce.”
""")

st.sidebar.markdown("---")

st.sidebar.markdown("""
### ⚠️ Limitations  
- Always check answer with sources— the linked documents before sharing with a client.  
- Answers are limited to materials in the RFP repository.  
- No predictions—the bot summarizes existing info; it does not forecast future targets.  
- Confidentiality—**do not paste client-sensitive content** into the chat.
""")

# ====== Constants ======
BASE_DOC_URL = "https://yourdomain.com/docs/"  # Update to your hosting location
SHEET_NAME = "ESGenieChatLogs"

# ====== Session State ======
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "query" not in st.session_state:
    st.session_state.query = ""
if "last_submitted_query" not in st.session_state:
    st.session_state.last_submitted_query = ""
if "run_query" not in st.session_state:
    st.session_state.run_query = False

# ====== Hero Section ======
col1, col2 = st.columns([1, 8])
with col1:
    st.image("logo.png", width=400)
with col2:
    st.markdown("""
        <div style="display: flex; flex-direction: column; justify-content: center; height: 100%;">
            <h1 style="margin-bottom: 0;">RFP Assistant</h1>
            <p style="font-size: 16px; margin-top: 4px;">
                Generate responses with citations and document links
            </p>
        </div>
    """, unsafe_allow_html=True)
st.markdown("")

# ====== Suggestion Buttons ======
col0, col1, col2, col3, col4, col5 = st.columns([1.5, 2.7, 3, 3.2, .3, .5], gap="small")
with col0:
    st.markdown("#### Try a prompt:")
with col1:
    if st.button("What are Bain sustainability commitments?", key="prompt1"):
        st.session_state.query = "What are Bain's sustainability commitments?"
with col2:
    if st.button("What are Bain's waste diversion efforts/policy?", key="prompt2"):
        st.session_state.query = "What are Bain's waste diversion policy?"
with col3:
    if st.button("What are Bain's ISO certifications, if any?", key="prompt3"):
        st.session_state.query = "What are Bain's ISO certifications, if any?"

# Handle button behavior
if st.session_state.get("prompt1"):
    st.session_state.query = "What are Bain's sustainability commitments?"
elif st.session_state.get("prompt2"):
    st.session_state.query = "What are Bain's waste diversion policy?"
elif st.session_state.get("prompt3"):
    st.session_state.query = "What are Bain's ISO certifications, if any?"

# Add vertical space after prompt buttons
st.markdown("<div style='margin-top: 30px;'></div>", unsafe_allow_html=True)

# ====== Configuration ======
openai.api_key = st.secrets["openai_api_key"]
EMBEDDING_MODEL = "text-embedding-3-small"
CHAT_MODEL = "gpt-4"
TOP_K = 20
MAX_CONTEXT_TOKENS = 6000

# ====== Document Priorities ======
DOCUMENT_PRIORITIES = {
    "GRI report 2023.pdf": 1,
    "Environmental-policy.pdf": 2,
    "climate-transition-plan-2024.pdf": 3,
    "Code of Conduct - client version.pdf": 4,
    "Bain DEI report.pdf": 5,
    "2023_carbon-credit-disclosure.pdf": 6,
    "bain-wef-and-tcfd-report-2023.pdf": 7,
    "Bain Overview FAQ.pdf": 8,
    "Full RFP FAQ Export.pdf":9,
    "Client RFP deck.pdf": 10,
    "Global Safety & Security FAQ.pdf": 11,
    "bain-sustainable-procurement-policy.pdf": 12,
    "Social Impact.pdf": 13,
    "human-rights-statement-05.2024.pdf": 14,
    "sustainable-procurement-factsheet-v3-05.02.2024.pdf": 15,
    "Diversity and Inclusion FAQ.pdf": 16
}

# ====== Google Sheets Logger ======
def log_chat_to_gsheet(question, answer, sources):
    try:
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
        creds = ServiceAccountCredentials.from_json_keyfile_name("service_account.json", scope)
        client = gspread.authorize(creds)
        sheet = client.open(SHEET_NAME).sheet1

        sources_text = ", ".join([
            f"{chunk['metadata'].get('document', 'Unknown')} (page {chunk['metadata'].get('page', '?')})"
            for chunk in sources
        ])
        sheet.append_row([
            pd.Timestamp.now().isoformat(),
            question,
            answer or "No answer generated.",
            sources_text
        ])
    except Exception as e:
        st.error(f"Logging to Google Sheets failed: {e}")

# ====== Load Vector DB & Metadata ======
@st.cache_resource
def load_resources():
    index = faiss.read_index("4.my_vector_db.index")
    with open("5.metadata.json", "r", encoding="utf-8") as f:
        metadata = json.load(f)
    return index, metadata

# ====== Embedding Function ======
def get_embedding(text):
    try:
        response = openai.embeddings.create(input=[text], model=EMBEDDING_MODEL)
        return np.array(response.data[0].embedding, dtype="float32")
    except Exception as e:
        st.error(f"Embedding error: {e}")
        return None

# ====== Chunk Search ======
def search_chunks(query_text, index, metadata):
    ids = [item["id"] for item in metadata]
    texts = [item["text"] for item in metadata]
    meta_lookup = {item["id"]: item["metadata"] for item in metadata}

    query_vec = get_embedding(query_text)
    if query_vec is None:
        return []

    D, I = index.search(np.array([query_vec]), TOP_K)
    results = []

    for rank, i in enumerate(I[0]):
        if i == -1:
            continue
        chunk_id = ids[i]
        text = texts[i]
        meta = meta_lookup.get(chunk_id, {})
        doc = meta.get("document", "Unknown")
        page = meta.get("page", "?")
        priority = DOCUMENT_PRIORITIES.get(doc, 1000)
        score = float(D[0][rank]) * (1 + priority / 100)
        results.append({
            "text": text,
            "score": score,
            "metadata": meta
        })

    return sorted(results, key=lambda x: x["score"])

# ====== GPT Answer Generator with Inline Source Tags ======
def generate_answer(query, context_chunks):
    encoding = tiktoken.encoding_for_model("gpt-4")
    context_parts = []
    total_tokens = 0

    for chunk in context_chunks:
        doc = chunk["metadata"].get("document", "Unknown")
        page = chunk["metadata"].get("page", "?")
        tagged = f"{chunk['text']}\n(Source: {doc}, page {page})"
        token_count = len(encoding.encode(tagged))
        if total_tokens + token_count > MAX_CONTEXT_TOKENS:
            break
        context_parts.append(tagged)
        total_tokens += token_count

    context_text = "\n\n".join(context_parts)

    try:
        response = openai.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {"role": "system", "content": (
                    "You are an expert assistant. Only answer using the provided context."
                    "Every claim MUST be followed by an inline citation in the format (Source: filename, page X). Do not skip citations."
                )},
                {"role": "user", "content": f"{context_text}\n\nQuestion: {query}"}
            ],
            max_tokens=1000
        )
        return response.choices[0].message.content.strip(), context_chunks
    except Exception as e:
        st.error(f"OpenAI error: {e}")
        return None, []

# ====== Export Chat History as Word Document ======
def generate_docx(chat_history):
    doc = Document()

    # Add title with current date
    today_str = datetime.now().strftime("%B %d, %Y")  # Example: June 11, 2025
    doc.add_heading(f"Chat Export – {today_str}", level=1)

    if not chat_history:
        doc.add_paragraph("No chat history available.")
    else:
        for i, chat in enumerate(chat_history, 1):
            doc.add_heading(f"Question {i}", level=2)
            doc.add_paragraph(chat['question'])
            doc.add_heading("Answer", level=3)
            doc.add_paragraph(chat['answer'])
            doc.add_paragraph("")  # spacing

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ====== User Input with Perfectly Aligned Submit Button ======
st.markdown("### 🔍 Ask your question:")

# Layout: wider input, narrow button
input_col, button_col = st.columns([8, 1])

with input_col:
    query = st.text_area(
        label="",
        value=st.session_state.get("query", ""),
        key="query_input",
        height=68,
        placeholder="Type your ESG question here..."
    )

# Inject vertical alignment fix
with button_col:
    st.markdown(
        """<div style="margin-top: 45px;">""", unsafe_allow_html=True
    )
    submit_clicked = st.button("Submit", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

# Trigger processing
if submit_clicked and query.strip():
    st.session_state.query = query
    st.session_state.run_query = True


# Run only on click
if submit_clicked and query.strip():
    st.session_state.query = query
    st.session_state.run_query = True


if st.session_state.run_query:
    with st.spinner("Processing..."):
        index, metadata = load_resources()
        top_chunks = search_chunks(query, index, metadata)
        answer, used_chunks = generate_answer(query, top_chunks) if top_chunks else ("No relevant context found.", [])
        st.session_state.chat_history.append({
            "question": query,
            "answer": answer,
            "sources": used_chunks
        })
        log_chat_to_gsheet(query, answer, used_chunks)
    st.session_state.last_submitted_query = query
    st.session_state.run_query = False  # ✅ Reset trigger

# Step 3: If new query has been submitted, run it
if st.session_state.get("run_query", False):
    with st.spinner("Processing..."):
        index, metadata = load_resources()
        top_chunks = search_chunks(query, index, metadata)
        answer, used_chunks = generate_answer(query, top_chunks) if top_chunks else ("No relevant context found.", [])

        st.session_state.chat_history.append({
            "question": query,
            "answer": answer,
            "sources": used_chunks
        })
        log_chat_to_gsheet(query, answer, used_chunks)

    st.session_state.run_query = False  # ✅ Mark query processed


# ====== Trigger Query Processing If New ======
if query.strip() and (query != st.session_state.get("last_submitted_query", "")):
    st.session_state.last_submitted_query = query
    with st.spinner("Processing..."):
        index, metadata = load_resources()
        top_chunks = search_chunks(query, index, metadata)
        answer, used_chunks = generate_answer(query, top_chunks) if top_chunks else ("No relevant context found.", [])
        st.session_state.chat_history.append({
            "question": query,
            "answer": answer,
            "sources": used_chunks
        })
        log_chat_to_gsheet(query, answer, used_chunks)

# ====== Clear & Export Buttons Side-by-Side ======
col1, col2 = st.columns([1, 4.6])

with col1:
    if st.button("🗑️ Clear conversation History", key="clear_chat_button"):
        # Clear all relevant session state
        st.session_state.clear_trigger = True

with col2:
    docx_file = generate_docx(st.session_state.chat_history)
    st.download_button(
        label="📥 Export",
        data=docx_file,
        file_name="Chat_Export.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="export_docx"
    )

# ✅ SAFELY PERFORM CLEARING + RERUN AFTER FIRST RENDER CYCLE
if st.session_state.get("clear_trigger", False):
    st.session_state.clear_trigger = False
    st.session_state.chat_history.clear()
    st.session_state.query = ""
    st.session_state.last_submitted_query = ""
    st.session_state.run_query = False

    # Optional toast confirmation
    st.success("Conversation history cleared.")


# ====== Display History ======
for i, chat in enumerate(reversed(st.session_state.chat_history), 1):
    st.markdown(f"#### Question {len(st.session_state.chat_history) - i + 1}")
    st.markdown(chat["question"])

    # Display generated answer
    st.markdown("#### Answer")
    inline = re.sub(r'\(Source: (.*?), page (\d+)\)', r'📝 [\1 – p.\2]', chat["answer"])
    st.markdown(inline, unsafe_allow_html=True)

    # Show only verbatim chunks that were actually cited, in citation order
    if chat["sources"]:
        # Get ordered list of (document, page) from answer text
        ordered_refs = re.findall(r'\(Source: (.*?), page (\d+)\)', chat["answer"])

        # Build lookup for quick access to matching chunks
        ref_to_chunk = {
            (chunk["metadata"].get("document", "Unknown"), str(chunk["metadata"].get("page", "?"))): chunk
            for chunk in chat["sources"]
        }

        # Keep only chunks that are actually cited, in citation order (no duplicates)
        seen = set()
        ordered_chunks = []
        for ref in ordered_refs:
            if ref not in seen and ref in ref_to_chunk:
                ordered_chunks.append(ref_to_chunk[ref])
                seen.add(ref)

        if ordered_chunks:
            with st.expander("📄 Show Verbatim Source Excerpts"):
                for idx, chunk in enumerate(ordered_chunks, 1):
                    doc = chunk["metadata"].get("document", "Unknown")
                    page = chunk["metadata"].get("page", "?")
                    st.markdown(f"**{idx}. {doc}, page {page}**")
                   
                    # Step 1: Normalize and extract bullets
                    raw_text = chunk["text"]

                    # Normalize whitespace
                    text = re.sub(r'[ \t]+', ' ', raw_text)

                    # Replace all bullet-like characters (• ● * -) with a consistent marker
                    # Capture only bullets that are at the start of lines or mid-sentence
                    text = re.sub(r'(^|\n)[ \t]*[-•●*][ \t]+', r'\1|||BULLET|||', text)
                    text = re.sub(r'[•●*]', '|||BULLET|||', text)  # Catch inline ones

                    # Split and process bullet parts
                    parts = text.split('|||BULLET|||')
                    formatted_parts = []

                    for i, part in enumerate(parts):
                        part = part.strip()
                        if not part:
                            continue

                        # Remove bullets that start with lowercase or conjunctions
                        if re.match(r'^(and|or|but|so|to)\b', part, flags=re.IGNORECASE):
                            continue
                        if not re.match(r'^[A-Z0-9]', part):
                            continue

                        if i == 0:
                            formatted_parts.append(part)
                        else:
                            formatted_parts.append(f"- {part}")

                    # Join and enforce spacing for Markdown rendering
                    cleaned_text = "\n\n".join(formatted_parts).strip()

                    # Display clean bullet structure
                    st.markdown(cleaned_text)
                    
    # Display source document links
    st.markdown("#### 📚 Additional Sources")
    grouped = defaultdict(list)
    for chunk in chat["sources"]:
        meta = chunk["metadata"]
        doc = meta.get("document", "Unknown")
        page = meta.get("page", "?")
        grouped[doc].append(page)

    for doc in grouped.keys():
        url = next(
            (chunk["metadata"].get("link", "#") for chunk in chat["sources"]
             if chunk["metadata"].get("document") == doc),
            "#"
        )
        st.markdown(f"- [**{doc}**]({url})")
